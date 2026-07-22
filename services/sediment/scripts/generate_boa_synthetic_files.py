"""Generate strict BOA-style synthetic NAS files."""
from __future__ import annotations

import json
import random
import sys
from datetime import date, timedelta
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from scripts.boa_lib import (
    BOA_MANIFEST_PATH,
    BOA_NAS_ROOT,
    BoaError,
    assert_no_forbidden_substitute,
    choose_docs_for_profile,
    detect_pii_text,
    load_taxonomy,
    mkdirs,
    parse_profile_arg,
    remove_generated_roots_for_fresh_run,
    repo_rel,
    reset_generated_roots,
    sha256_file,
    write_json,
    write_report,
)

INSURERS = ["ABL", "AIA", "AXA", "DB", "KB", "MG", "NH", "교보", "농협", "동양", "라이나", "롯데", "메리츠", "미래에셋", "삼성생명"]
ROLES = ["데스크", "상담실", "보험담당", "총괄실장", "교정팀", "기공실", "경영지원"]
SYN_PATIENTS = [f"BOA-P{n:04d}" for n in range(1, 401)]


def _safe_name(s: str) -> str:
    return s.replace("/", "_").replace("\\", "_").replace(":", "_")


def _pick_ext(dt: dict[str, Any], global_remaining: dict[str, int], rng: random.Random) -> str:
    allowed = list(dt.get("extensions") or ["docx"])
    normalized = ["png_jpg" if x in {"png", "jpg", "jpeg"} else x for x in allowed]
    viable = [x for x in normalized if global_remaining.get(x, 0) > 0]
    if not viable:
        picked = normalized[0]
    else:
        picked = max(viable, key=lambda x: global_remaining[x])
        global_remaining[picked] -= 1
    if picked == "png_jpg":
        return rng.choice(["png", "jpg"])
    return picked


def _title(area: dict[str, Any], dt: dict[str, Any], idx: int, rng: random.Random) -> str:
    examples = dt.get("filename_examples") or []
    pattern = dt.get("filename_pattern")
    if pattern:
        insurer = INSURERS[(idx - 1) % len(INSURERS)]
        return pattern.format(insurer=insurer)
    if examples:
        base = Path(examples[(idx - 1) % len(examples)]).stem
    else:
        base = f"{area['id']}_{dt['id']}_{idx:03d}"
    suffixes = ["", "_수정", "_최종", "_v2", "_2026"]
    return _safe_name(base + suffixes[idx % len(suffixes)])


def _folder(area: dict[str, Any], dt: dict[str, Any]) -> Path:
    return BOA_NAS_ROOT / (dt.get("folder_override") or area["folder"])


def _body(
    title: str,
    area: dict[str, Any],
    dt: dict[str, Any],
    idx: int,
    rng: random.Random,
    scenario_terms: dict[str, list[str]],
) -> str:
    patient = SYN_PATIENTS[(idx * 17) % len(SYN_PATIENTS)]
    insurer = INSURERS[(idx - 1) % len(INSURERS)]
    sections = dt.get("required_sections") or ["목적", "체크리스트", "담당자 확인", "주의사항"]
    modules = area.get("lecture_modules") or []
    terms = sorted({term for module in modules for term in scenario_terms.get(module, [])})
    lines = [
        title,
        "",
        f"- 문서유형: {dt['id']}",
        f"- 업무영역: {area['id']}",
        f"- 담당역할: {ROLES[idx % len(ROLES)]}",
        f"- 합성 케이스 ID: {patient}",
        f"- 보험사 예시: {insurer}",
        "- 실제 개인정보 포함 여부: false",
        f"- 강의 핵심어: {', '.join(terms)}",
        "",
    ]
    for sidx, section in enumerate(sections, start=1):
        lines += [
            f"{sidx}. {section}",
            f"{section} 단계에서는 데스크 확인, 상담 기록, 차트 확인, 서류 누락 방지를 함께 점검한다.",
            f"{title} 문서는 보아치과 강의용 synthetic 자료이며 {patient} 케이스와 연결된다.",
            "체크 항목: 환자 안내, 서명 여부, 진료일, 치아번호, 처치명, 비용, 담당자 확인.",
            "",
        ]
    lines += [
        "강의 질문 연결",
        "이 문서는 Sediment 검색과 ask 시나리오에서 근거 문서로 사용된다.",
        "핵심 용어: " + ", ".join(terms),
        "답변은 반드시 이 synthetic 문서의 범위 안에서만 구성해야 한다.",
    ]
    text = "\n".join(lines)
    hits = detect_pii_text(text)
    if hits:
        raise BoaError(f"generated PII-like text in {title}: {hits}")
    return text


def _write_docx(path: Path, title: str, body: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body.splitlines()[1:]:
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(path)


def _write_xlsx(path: Path, title: str, body: str, rng: random.Random) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "체크리스트"
    ws.append(["문서명", title])
    ws.append(["실제 개인정보 포함", "false"])
    ws.append([])
    ws.append(["합성케이스ID", "진료일", "구분", "치아번호", "처치명", "서류상태", "담당자", "비고"])
    treatments = ["스케일링", "레진", "크라운", "임플란트 상담", "교정 상담", "유지장치", "치과치료확인서"]
    statuses = ["준비", "확인완료", "서명대기", "제출대기", "보완필요"]
    start = date(2026, 1, 3)
    for i in range(18 + rng.randint(0, 18)):
        ws.append([
            SYN_PATIENTS[(i * 13) % len(SYN_PATIENTS)],
            (start + timedelta(days=i * 3)).isoformat(),
            "synthetic",
            rng.choice(["16", "21", "26", "36", "46", "전악"]),
            rng.choice(treatments),
            rng.choice(statuses),
            rng.choice(ROLES),
            "강의용 synthetic row",
        ])
    wb.save(path)


def _write_pdf(path: Path, title: str, body: str) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    _, height = A4
    y = height - 50
    c.setFont("Helvetica", 13)
    c.drawString(50, y, title[:80])
    y -= 28
    c.setFont("Helvetica", 9)
    for line in body.replace("실제 개인정보", "synthetic data").splitlines():
        if not line.strip():
            y -= 8
            continue
        c.drawString(50, y, line[:105])
        y -= 14
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 9)
            y = height - 50
    c.save()


def _write_image(path: Path, title: str) -> None:
    img = Image.new("RGB", (900, 520), (245, 247, 250))
    d = ImageDraw.Draw(img)
    d.rectangle((30, 30, 870, 490), outline=(40, 90, 120), width=3)
    d.text((60, 70), "BOA Dental Synthetic Asset", fill=(20, 40, 60))
    d.text((60, 120), title[:70], fill=(20, 40, 60))
    d.text((60, 180), "contains_real_phi=false", fill=(80, 80, 80))
    img.save(path)


def _write_md_json(path: Path, title: str, body: str) -> None:
    if path.suffix == ".json":
        path.write_text(json.dumps({"title": title, "body": body, "synthetic": True}, ensure_ascii=False, indent=2), encoding="utf-8")
    else:
        path.write_text(body, encoding="utf-8")


def main() -> None:
    ap = parse_profile_arg()
    args = ap.parse_args()
    taxonomy = load_taxonomy()
    mkdirs()
    if BOA_MANIFEST_PATH.exists():
        reset_generated_roots()
    else:
        remove_generated_roots_for_fresh_run()

    rng = random.Random(int(taxonomy["execution_policy"]["deterministic_seed"]))
    scenario_terms = {s["id"]: list(s.get("expected_terms") or []) for s in taxonomy["golden_scenarios"]}
    remaining = dict(taxonomy["file_type_mix"])
    scale = sum(remaining.values()) / int(taxonomy["target_profile"]["total_files"])
    if args.profile != "lecture_360":
        target = int(taxonomy["profiles"][args.profile]["total_files"])
        base_total = int(taxonomy["target_profile"]["total_files"])
        remaining = {k: round(v * target / base_total) for k, v in remaining.items()}
        diff = target - sum(remaining.values())
        keys = sorted(remaining, key=remaining.get, reverse=True)
        for i in range(abs(diff)):
            remaining[keys[i % len(keys)]] += 1 if diff > 0 else -1
    del scale

    docs = choose_docs_for_profile(taxonomy, args.profile)
    manifest_files: list[dict[str, Any]] = []
    for n, spec in enumerate(docs, start=1):
        area = spec["area"]
        dt = spec["doc_type"]
        ext = _pick_ext(dt, remaining, rng)
        title = _title(area, dt, n, rng)
        if not title.endswith(f".{ext}"):
            filename = f"{n:04d}-{title}.{ext}"
        else:
            filename = f"{n:04d}-{title}"
        folder = _folder(area, dt)
        folder.mkdir(parents=True, exist_ok=True)
        path = folder / filename
        assert_no_forbidden_substitute(path)
        body = _body(Path(filename).stem, area, dt, n, rng, scenario_terms)
        if ext == "docx":
            _write_docx(path, Path(filename).stem, body)
        elif ext == "xlsx":
            _write_xlsx(path, Path(filename).stem, body, rng)
        elif ext == "pdf":
            _write_pdf(path, Path(filename).stem, body)
        elif ext in {"png", "jpg"}:
            _write_image(path, Path(filename).stem)
        elif ext in {"md", "json"}:
            _write_md_json(path, Path(filename).stem, body)
        else:
            raise BoaError(f"unsupported extension: {ext}")
        meta = {
            "tenant_slug": taxonomy["tenant_slug"],
            "source_pattern": "boa_screenshot",
            "folder_area": area["id"],
            "doc_type": dt["id"],
            "year": 2026 if "2026" in filename or area["id"] == "staff_2026" else 2025,
            "owner_role": ROLES[n % len(ROLES)],
            "lecture_module": (area.get("lecture_modules") or ["nas_orientation"])[0],
            "demo_relevance": "high" if n % 3 == 0 else "normal",
            "contains_real_phi": False,
            "synthetic_patient_ids": [SYN_PATIENTS[(n * 17) % len(SYN_PATIENTS)]],
            "expected_queries": [s["anchor_query"] for s in taxonomy["golden_scenarios"] if s["id"] in area.get("lecture_modules", [])],
            "title": Path(filename).stem,
            "body_text": body,
            "file_extension": ext,
        }
        meta_path = path.with_name(path.name + ".metadata.json")
        write_json(meta_path, meta)
        item = {
            "path": repo_rel(path),
            "metadata_path": repo_rel(meta_path),
            "sha256": sha256_file(path),
            "metadata_sha256": sha256_file(meta_path),
            "doc_type": dt["id"],
            "area": area["id"],
            "extension": ext,
        }
        manifest_files.append(item)

    manifest = {
        "profile": args.profile,
        "taxonomy": repo_rel(Path("services/sediment/data/boa_file_taxonomy.yaml")),
        "count": len(manifest_files),
        "files": manifest_files,
    }
    write_json(BOA_MANIFEST_PATH, manifest)
    report = {
        "ok": True,
        "profile": args.profile,
        "files": len(manifest_files),
        "remaining_file_type_budget": remaining,
        "manifest": repo_rel(BOA_MANIFEST_PATH),
    }
    write_report("generate", report)
    print(json.dumps(report, ensure_ascii=False, sort_keys=True))


if __name__ == "__main__":
    main()
